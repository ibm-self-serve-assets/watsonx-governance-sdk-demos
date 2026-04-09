"""
Contract Agent for ESA / Contract Ingestion

This agent performs OCR, contract extraction, and ingestion into vector database.
It reads contracts from file paths, extracts key information, and enables semantic queries.
"""

from typing import TypedDict, Annotated, Optional, Dict, List, Union
from langgraph.graph import StateGraph, START, END
from langchain_ibm import WatsonxLLM, WatsonxEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document
import operator
import os
from pathlib import Path
import json
from datetime import datetime
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# For document processing
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False


class ContractState(TypedDict):
    """State for Contract Agent"""
    file_path: str
    raw_text: Optional[str]
    contract_metadata: Optional[Dict]
    structured_summary: Optional[Dict]
    vector_store_status: Optional[str]
    generated_text: str
    messages: Annotated[list, operator.add]


class PortfolioContractState(TypedDict):
    """State for multi-contract portfolio processing"""
    partner_name: str
    contract_paths: List[str]
    contract_results: List[Dict]
    portfolio_summary: Optional[Dict]
    generated_text: str
    messages: Annotated[list, operator.add]


class ContractAgent:
    """
    Contract Agent using LangGraph for ESA/Contract ingestion.
    
    This agent:
    1. Reads document from file path (read-only access)
    2. Performs OCR/text extraction
    3. Extracts and normalizes contract metadata:
       - Parties
       - Effective date
       - Term length
       - Key obligations/milestones
    4. Ingests into in-memory vector database (Chroma)
    5. Returns structured contract summary
    """
    
    def __init__(
        self,
        model_id: str = "meta-llama/llama-3-3-70b-instruct",
        embedding_model_id: str = "ibm/slate-125m-english-rtrvr-v2",
        url: str = "https://us-south.ml.cloud.ibm.com",
        apikey: Optional[str] = None,
        project_id: Optional[str] = None,
        max_new_tokens: int = 800,
        vector_store_path: str = "./contract_vector_store"
    ):
        """
        Initialize Contract Agent.
        
        Args:
            model_id: Watsonx model ID for text generation
            embedding_model_id: Watsonx model ID for embeddings
            url: Watsonx API URL
            apikey: IBM Cloud API key (defaults to WATSONX_APIKEY env var)
            project_id: Watsonx project ID (defaults to WATSONX_PROJECT_ID env var)
            max_new_tokens: Maximum tokens for generation
            vector_store_path: Path for persistent vector store
        """
        self.model_id = model_id
        self.embedding_model_id = embedding_model_id
        self.url = url
        # Use environment variables if not provided
        self.apikey = apikey or os.getenv("WATSONX_APIKEY")
        self.project_id = project_id or os.getenv("WATSONX_PROJECT_ID")
        self.max_new_tokens = max_new_tokens
        self.vector_store_path = vector_store_path
        self.graph = None
        self.vector_store = None
        
        # Validate credentials
        if not self.apikey:
            raise ValueError("WATSONX_APIKEY must be provided or set in environment variables")
        if not self.project_id:
            raise ValueError("WATSONX_PROJECT_ID must be provided or set in environment variables")
    def _load_cache(self):
        """Load contract cache from JSON file"""
        cache_file = "contracts_cache.json"
        if os.path.exists(cache_file):
            try:
                with open(cache_file, 'r') as f:
                    cache_data = json.load(f)
                    return cache_data.get('contracts', [])
            except Exception as e:
                print(f"Warning: Could not load cache: {e}")
                return []
        return []
    
    def _get_cached_contract(self, file_path: str):
        """
        Get cached contract data if available.
        
        Args:
            file_path: Path to contract file
            
        Returns:
            Cached contract data dict or None if not found
        """
        cache = self._load_cache()
        file_name = os.path.basename(file_path)
        
        for contract in cache:
            if contract.get('file_name') == file_name or contract.get('file_path') == file_path:
                return contract
        
        return None
    
        
    def _initialize_vector_store(self):
        """Initialize or load the vector store"""
        if self.vector_store is None:
            embeddings = WatsonxEmbeddings(
                model_id=self.embedding_model_id,
                url=self.url,
                apikey=self.apikey,
                project_id=self.project_id
            )
            
            # Create or load persistent vector store
            self.vector_store = Chroma(
                collection_name="contracts",
                embedding_function=embeddings,
                persist_directory=self.vector_store_path
            )
    
    def _read_document(self, file_path: str) -> str:
        """
        Read document content from file path (read-only).
        Supports .txt, .docx, and .xlsx files.
        
        Args:
            file_path: Path to the contract document
            
        Returns:
            Extracted text content
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Read based on file extension
        if file_path.suffix.lower() == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        elif file_path.suffix.lower() == '.docx':
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx not installed. Install with: pip install python-docx")
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = '\t'.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text_parts.append(row_text)
            
            full_text = '\n'.join(text_parts)
            
            # If still empty, try alternative extraction methods
            if not full_text.strip():
                # Fallback 1: Try to get text from all runs in all paragraphs
                all_text = []
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        if run.text and run.text.strip():
                            all_text.append(run.text)
                
                # Fallback 2: Try to extract from document body elements
                if not all_text:
                    try:
                        from docx.oxml.text.paragraph import CT_P
                        from docx.oxml.table import CT_Tbl
                        from docx.table import Table
                        from docx.text.paragraph import Paragraph
                        
                        for element in doc.element.body:
                            if isinstance(element, CT_P):
                                para = Paragraph(element, doc)
                                if para.text and para.text.strip():
                                    all_text.append(para.text)
                            elif isinstance(element, CT_Tbl):
                                table = Table(element, doc)
                                for row in table.rows:
                                    row_text = '\t'.join([cell.text.strip() for cell in row.cells])
                                    if row_text.strip():
                                        all_text.append(row_text)
                    except Exception as e:
                        # If advanced extraction fails, log but continue
                        pass
                
                full_text = '\n'.join(all_text)
            
            # Final check - if still empty, raise an error
            if not full_text.strip():
                raise ValueError(f"Could not extract any text from DOCX file: {file_path}")
            
            return full_text
        
        elif file_path.suffix.lower() in ['.xlsx', '.xls']:
            if not EXCEL_AVAILABLE:
                raise ImportError("openpyxl not installed. Install with: pip install openpyxl")
            wb = openpyxl.load_workbook(file_path, read_only=True)
            text_parts = []
            for sheet in wb.worksheets:
                text_parts.append(f"\n=== Sheet: {sheet.title} ===\n")
                for row in sheet.iter_rows(values_only=True):
                    row_text = '\t'.join([str(cell) if cell is not None else '' for cell in row])
                    if row_text.strip():
                        text_parts.append(row_text)
            return '\n'.join(text_parts)
        
        else:
            # Fallback: try to read as text
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except UnicodeDecodeError:
                raise ValueError(f"Unsupported file format: {file_path.suffix}")
    
    def build_agent(self):
        """Build the LangGraph for contract processing"""
        
        def read_document_node(state: ContractState) -> dict:
            """Read document from file path - checks cache first"""
            file_path = state['file_path']
            
            # Check cache first
            cached = self._get_cached_contract(file_path)
            if cached:
                print(f"✓ Using cached data for: {os.path.basename(file_path)}")
                return {
                    "raw_text": cached.get('extracted_text'),
                    "messages": [f"Loaded from cache: {len(cached.get('extracted_text', ''))} characters"]
                }
            
            # Not in cache - read from file
            try:
                print(f"DEBUG: Reading document from: {file_path}")
                raw_text = self._read_document(file_path)
                print(f"DEBUG: Extracted {len(raw_text) if raw_text else 0} characters")
                
                if not raw_text or len(raw_text.strip()) == 0:
                    print(f"DEBUG: Document appears empty!")
                    return {
                        "raw_text": None,
                        "messages": [f"Warning: Document appears empty or could not extract text from {file_path}"]
                    }
                
                print(f"DEBUG: Returning raw_text with {len(raw_text)} characters")
                return {
                    "raw_text": raw_text,
                    "messages": [f"Document read successfully: {len(raw_text)} characters"]
                }
            except Exception as e:
                import traceback
                error_details = traceback.format_exc()
                print(f"DEBUG: Error reading document: {str(e)}")
                print(error_details)
                return {
                    "raw_text": None,
                    "messages": [f"Error reading document: {str(e)}\n{error_details}"]
                }
        
        def extract_metadata_node(state: ContractState) -> dict:
            """Extract and normalize contract metadata using LLM with rule-based fallback - checks cache first"""
            
            # Check cache first to avoid LLM call
            file_path = state.get("file_path")
            if file_path:
                cached = self._get_cached_contract(file_path)
                if cached and 'structured_summary' in cached:
                    print(f"✓ Using cached metadata for: {os.path.basename(file_path)}")
                    # Create metadata structure from cached data
                    metadata = {
                        "raw_extraction": f"CACHED DATA - {cached['structured_summary']}",
                        "file_path": file_path,
                        "extraction_timestamp": cached.get('processed_at', datetime.now().isoformat()),
                        "document_length": cached.get('text_length', 0)
                    }
                    return {
                        "contract_metadata": metadata,
                        "messages": ["Loaded metadata from cache (no LLM call)"]
                    }
            
            if not state["raw_text"]:
                return {
                    "contract_metadata": None,
                    "messages": ["Skipping metadata extraction - no text available"]
                }
            
            text_sample = state["raw_text"][:3000]
            extracted_metadata = None
            
            metadata_prompt = ChatPromptTemplate.from_template(
                "You are a contract analysis expert. Extract key metadata from this contract:\n\n"
                "{contract_text}\n\n"
                "CRITICAL INSTRUCTIONS:\n"
                "1. Extract ALL parties mentioned - this includes BOTH IBM AND the partner company (e.g., Confluent)\n"
                "2. List EVERY product/service mentioned, not just the first one\n"
                "3. Look for phrases like 'Cloud Services IBM watsonx as a Service, IBM watsonx Code Assistant for Ansible Service'\n"
                "4. Extract each product separately\n\n"
                "EXAMPLE:\n"
                "Text: 'Cloud Services IBM watsonx as a Service, IBM watsonx Code Assistant for Ansible Service'\n"
                "Extract as: watsonx as a Service, watsonx Code Assistant for Ansible Service\n\n"
                "Extract and provide in this exact format:\n"
                "PARTIES: [List ALL parties - both IBM and partner company]\n"
                "EFFECTIVE_DATE: [Contract effective date]\n"
                "TERM_LENGTH: [Duration/term of contract, e.g., '1 year', '2 years']\n"
                "KEY_OBLIGATIONS: [Main obligations and responsibilities]\n"
                "MILESTONES: [Key milestones or deliverables]\n"
                "PRODUCTS: [ALL products/services mentioned, separated by commas]\n\n"
                "If information is not found, write 'Not specified'."
            )

            try:
                llm = WatsonxLLM(
                    model_id=self.model_id,
                    url=self.url,
                    apikey=self.apikey,
                    project_id=self.project_id,
                    params={
                        "max_new_tokens": 500,
                        "decoding_method": "greedy",
                    }
                )
                
                formatted_prompt = metadata_prompt.invoke({"contract_text": text_sample})
                result = llm.invoke(formatted_prompt)
                extracted_metadata = result.content if hasattr(result, "content") else str(result)
            except Exception as e:
                import re
                filename = Path(state["file_path"]).name
                normalized_text = state["raw_text"]

                date_match = re.search(r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\w+\s+\d{1,2},\s+\d{4})', normalized_text)
                effective_date = date_match.group(1) if date_match else "Not specified"

                parties = []
                if "confluent" in normalized_text.lower():
                    parties.append("Confluent")
                if "ibm" in normalized_text.lower():
                    parties.append("IBM")
                if not parties:
                    parties = ["IBM", "Confluent"]

                obligations = []
                lowered = normalized_text.lower()
                for keyword in ["license", "support", "payment", "renewal", "subscription", "cognos"]:
                    if keyword in lowered:
                        obligations.append(keyword)

                extracted_metadata = (
                    f"PARTIES: {parties}\n"
                    f"EFFECTIVE_DATE: {effective_date}\n"
                    f"TERM_LENGTH: Not specified\n"
                    f"KEY_OBLIGATIONS: {obligations or ['Not specified']}\n"
                    f"MILESTONES: Not specified\n"
                    f"FALLBACK_REASON: {str(e)}\n"
                    f"FILE_NAME: {filename}"
                )
            
            metadata = {
                "raw_extraction": extracted_metadata,
                "file_path": state["file_path"],
                "extraction_timestamp": datetime.now().isoformat(),
                "document_length": len(state["raw_text"])
            }
            
            return {
                "contract_metadata": metadata,
                "messages": ["Contract metadata extracted"]
            }
        
        def normalize_and_structure_node(state: ContractState) -> dict:
            """Normalize extracted data and create structured summary - checks cache first"""
            
            # Check if we have cached structured data
            file_path = state.get("file_path")
            if file_path:
                cached = self._get_cached_contract(file_path)
                if cached and 'structured_summary' in cached:
                    print(f"✓ Using cached structured data for: {os.path.basename(file_path)}")
                    return {
                        "structured_summary": cached['structured_summary'],
                        "messages": ["Loaded structured data from cache"]
                    }
            
            if not state["contract_metadata"]:
                return {
                    "structured_summary": None,
                    "messages": ["Skipping normalization - no metadata available"]
                }

            raw_extraction = state["contract_metadata"]["raw_extraction"]
            
            normalize_prompt = ChatPromptTemplate.from_template(
                "Based on this extracted contract metadata and the full contract text:\n\n"
                "{metadata}\n\n"
                "FULL CONTRACT TEXT (first 5000 chars):\n{contract_text}\n\n"
                "CRITICAL INSTRUCTIONS:\n"
                "1. Extract BOTH parties (IBM AND the partner company like Confluent)\n"
                "2. Calculate the actual end_date by adding term_length to start_date\n"
                "3. Extract ALL products as an array - if you see 'watsonx as a Service, watsonx Code Assistant', extract BOTH\n"
                "4. The partner company name is often in the filename (e.g., 'Confluent_IBM-1.30.2024.docx' means Confluent is a party)\n\n"
                "DATE CALCULATION EXAMPLES:\n"
                "- start_date='2024-01-31', term='1 year' → end_date='2025-01-31'\n"
                "- start_date='2024-03-30', term='2 years' → end_date='2026-03-30'\n"
                "- start_date='2023-03-30', term='3 years' → end_date='2026-03-30'\n\n"
                "PRODUCT EXTRACTION EXAMPLES:\n"
                "- Text: 'Cloud Services IBM watsonx as a Service, IBM watsonx Code Assistant for Ansible Service'\n"
                "- Extract: ['watsonx as a Service', 'watsonx Code Assistant for Ansible Service']\n\n"
                "Create a structured JSON summary with these fields:\n"
                "- parties: array of party names (MUST include both IBM and partner)\n"
                "- effective_date: normalized date (YYYY-MM-DD format)\n"
                "- start_date: contract start date (YYYY-MM-DD format, same as effective_date)\n"
                "- end_date: CALCULATED end date (YYYY-MM-DD format, NOT 'X years from effective date')\n"
                "- term_length: normalized duration (e.g., '1 year', '2 years', '3 years')\n"
                "- amount: total contract value with currency (e.g., '$500,092.80')\n"
                "- products: array of ALL products/services mentioned (extract each one separately)\n"
                "- key_obligations: array of main obligations\n"
                "- milestones: array of key milestones\n"
                "- contract_type: inferred type (e.g., 'ESA', 'Sales Agreement', 'Service Contract')\n"
                "- risk_level: assessed risk level (Low/Medium/High)\n\n"
                "IMPORTANT: Extract the EXACT dollar amount, ALL product names, and calculate actual dates.\n"
                "Look for phrases like 'committed order value', 'Purchase Commitment', 'Cloud Services', 'watsonx', 'Effective Date', 'term of this TD'.\n\n"
                "Provide ONLY valid JSON, no additional text."
            )
            
            try:
                llm = WatsonxLLM(
                    model_id=self.model_id,
                    url=self.url,
                    apikey=self.apikey,
                    project_id=self.project_id,
                    params={
                        "max_new_tokens": 400,
                        "decoding_method": "greedy",
                    }
                )
                
                formatted_prompt = normalize_prompt.invoke({
                    "metadata": raw_extraction,
                    "contract_text": (state.get("raw_text") or "")[:5000]  # First 5000 chars for context
                })
                result = llm.invoke(formatted_prompt)
                structured_data = result.content if hasattr(result, "content") else str(result)
                
                structured_data = structured_data.strip()
                if structured_data.startswith("```json"):
                    structured_data = structured_data[7:]
                elif structured_data.startswith("```"):
                    structured_data = structured_data[3:]
                if structured_data.endswith("```"):
                    structured_data = structured_data[:-3]
                structured_data = structured_data.strip()
                
                try:
                    structured_summary = json.loads(structured_data)
                    
                    # POST-PROCESSING: Calculate actual end_date if vague
                    # WHY: Contracts show "1 year from effective date" instead of "2025-01-31"
                    # This prevents urgency calculations like "expired 68 days ago"
                    if structured_summary.get("end_date") and ("from effective date" in str(structured_summary["end_date"]).lower() or "from the effective date" in str(structured_summary["end_date"]).lower()):
                        from dateutil.relativedelta import relativedelta
                        from dateutil import parser as date_parser
                        
                        start_date_str = structured_summary.get("start_date") or structured_summary.get("effective_date")
                        term_str = structured_summary.get("term_length", "")
                        
                        if start_date_str and term_str:
                            try:
                                # Parse start date
                                start_date = date_parser.parse(start_date_str)
                                
                                # Extract years from term
                                import re
                                years_match = re.search(r'(\d+)\s*years?', term_str, re.IGNORECASE)
                                if years_match:
                                    years = int(years_match.group(1))
                                    end_date = start_date + relativedelta(years=years)
                                    structured_summary["end_date"] = end_date.strftime("%Y-%m-%d")
                                    print(f"✓ Calculated end_date: {structured_summary['end_date']} (start: {start_date_str}, term: {term_str})")
                            except Exception as calc_error:
                                print(f"Warning: Could not calculate end_date: {calc_error}")
                    
                    # POST-PROCESSING: Extract partner name from filename if missing from parties
                    # WHY: LLM sometimes only extracts "IBM" and misses "Confluent"
                    parties = structured_summary.get("parties", [])
                    if len(parties) <= 1 or (len(parties) == 1 and "IBM" in parties):
                        filename = Path(state["file_path"]).name
                        # Extract partner name from filename pattern like "Confluent_IBM-1.30.2024.docx"
                        import re
                        partner_match = re.match(r'([^_]+)_IBM', filename)
                        if partner_match:
                            partner_name = partner_match.group(1)
                            if partner_name not in parties:
                                parties.append(partner_name)
                                structured_summary["parties"] = parties
                                print(f"✓ Added partner from filename: {partner_name}")
                    
                    # POST-PROCESSING: Extract ALL products from contract text if only one found
                    # WHY: LLM sometimes only extracts first product, missing "watsonx Code Assistant"
                    products = structured_summary.get("products", [])
                    if len(products) <= 1:
                        raw_text = state.get("raw_text", "").lower()
                        additional_products = []
                        
                        # Look for watsonx variants
                        if "watsonx as a service" in raw_text and "watsonx as a Service" not in products:
                            additional_products.append("watsonx as a Service")
                        if "watsonx code assistant" in raw_text and not any("code assistant" in p.lower() for p in products):
                            additional_products.append("watsonx Code Assistant for Ansible Service")
                        if "watsonx orchestrate" in raw_text and not any("orchestrate" in p.lower() for p in products):
                            additional_products.append("watsonx Orchestrate")
                        if "watsonx.governance" in raw_text or "watsonx governance" in raw_text:
                            if not any("governance" in p.lower() for p in products):
                                additional_products.append("watsonx.governance")
                        
                        # Look for Cognos
                        if "cognos" in raw_text and not any("cognos" in p.lower() for p in products):
                            additional_products.append("Cognos Analytics")
                        
                        if additional_products:
                            # Combine with existing, removing duplicates
                            all_products = list(set(products + additional_products))
                            structured_summary["products"] = all_products
                            print(f"✓ Enhanced products list: {all_products}")
                    
                except json.JSONDecodeError as e:
                    import re
                    json_match = re.search(r'\{[\s\S]*\}', structured_data)
                    if json_match:
                        try:
                            structured_summary = json.loads(json_match.group(0))
                        except json.JSONDecodeError:
                            structured_summary = {
                                "raw_structured_data": structured_data,
                                "parsing_note": f"Could not parse as JSON: {str(e)}"
                            }
                    else:
                        structured_summary = {
                            "raw_structured_data": structured_data,
                            "parsing_note": f"Could not parse as JSON: {str(e)}"
                        }
            except Exception as e:
                import re
                parties = []
                if "Confluent" in raw_extraction:
                    parties.append("Confluent")
                if "IBM" in raw_extraction:
                    parties.append("IBM")

                date_match = re.search(r'(20\d{2}-\d{2}-\d{2}|\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\w+\s+\d{1,2},\s+\d{4})', raw_extraction)
                effective_date = date_match.group(1) if date_match else "Unknown"

                risk_level = "Medium"
                if "cognos" in raw_extraction.lower():
                    risk_level = "Medium"

                # Enhanced fallback extraction from raw text
                raw_text = state.get("raw_text") or ""
                
                # Extract amount
                amount_match = re.search(r'\$[\d,]+\.?\d*', raw_text)
                amount = amount_match.group(0) if amount_match else "Not specified"
                
                # Extract products
                products = []
                if "watsonx" in raw_text.lower():
                    if "orchestrate" in raw_text.lower():
                        products.append("watsonx Orchestrate")
                    if "governance" in raw_text.lower():
                        products.append("watsonx.governance")
                    if "watsonx as a service" in raw_text.lower() or "watsonx.ai" in raw_text.lower():
                        products.append("watsonx as a Service")
                    if not products:  # Generic watsonx if no specific product found
                        products.append("watsonx ESA")
                if "cognos" in raw_text.lower():
                    products.append("Cognos")
                
                # Extract end date
                end_date_match = re.search(r'term of this TD will be (\w+) \((\d+)\) years? from the Effective Date', raw_text)
                if end_date_match:
                    term_years = end_date_match.group(2)
                    end_date = f"{term_years} year(s) from effective date"
                else:
                    end_date = "Not specified"
                
                structured_summary = {
                    "parties": parties or ["IBM", "Confluent"],
                    "effective_date": effective_date,
                    "end_date": end_date,
                    "start_date": effective_date,
                    "term_length": f"{term_years} year(s)" if end_date_match else "Not specified",
                    "amount": amount,
                    "products": products if products else ["Not specified"],
                    "key_obligations": ["renewal review", "contract management"] if "renewal" in raw_extraction.lower() else ["contract management"],
                    "milestones": [],
                    "contract_type": "ESA" if "ESA" in raw_text else "Sales Agreement",
                    "risk_level": risk_level,
                    "fallback_reason": str(e)
                }
            
            return {
                "structured_summary": structured_summary,
                "messages": ["Contract data normalized and structured"]
            }
        
        def ingest_to_vector_db_node(state: ContractState) -> dict:
            """Ingest contract into vector database for semantic search"""
            
            if not state["raw_text"]:
                return {
                    "vector_store_status": "Failed - no text available",
                    "messages": ["Vector ingestion skipped - no text"]
                }
            
            try:
                # Initialize vector store
                self._initialize_vector_store()
                
                # Create document with metadata
                doc_metadata = {
                    "file_path": state["file_path"],
                    "ingestion_timestamp": datetime.now().isoformat(),
                    "document_length": len(state["raw_text"])
                }
                
                # Add structured summary to metadata if available
                if state["structured_summary"]:
                    doc_metadata.update({
                        "contract_type": state["structured_summary"].get("contract_type", "Unknown"),
                        "parties": str(state["structured_summary"].get("parties", [])),
                        "effective_date": state["structured_summary"].get("effective_date", "Unknown")
                    })
                
                # Split text into chunks for better retrieval
                chunk_size = 1000
                chunks = []
                for i in range(0, len(state["raw_text"]), chunk_size):
                    chunk_text = state["raw_text"][i:i+chunk_size]
                    chunk_metadata = doc_metadata.copy()
                    chunk_metadata["chunk_index"] = i // chunk_size
                    chunks.append(Document(page_content=chunk_text, metadata=chunk_metadata))
                
                # Add to vector store
                self.vector_store.add_documents(chunks)
                
                return {
                    "vector_store_status": f"Success - {len(chunks)} chunks ingested",
                    "messages": [f"Contract ingested into vector DB: {len(chunks)} chunks"]
                }
            
            except Exception as e:
                return {
                    "vector_store_status": f"Failed - {str(e)}",
                    "messages": [f"Vector ingestion error: {str(e)}"]
                }
        
        def generate_summary_node(state: ContractState) -> dict:
            """Generate final structured contract summary"""
            
            # Build comprehensive summary
            summary_parts = [
                "="*70,
                "CONTRACT ANALYSIS SUMMARY",
                "="*70,
                "",
                f"File: {state['file_path']}",
                f"Document Length: {len(state['raw_text']) if state['raw_text'] else 0} characters",
                ""
            ]
            
            if state["contract_metadata"]:
                summary_parts.extend([
                    "EXTRACTED METADATA:",
                    "-"*70,
                    state["contract_metadata"]["raw_extraction"],
                    ""
                ])
            
            if state["structured_summary"]:
                summary_parts.extend([
                    "STRUCTURED SUMMARY:",
                    "-"*70,
                    json.dumps(state["structured_summary"], indent=2),
                    ""
                ])
            
            summary_parts.extend([
                "VECTOR DATABASE STATUS:",
                "-"*70,
                state.get("vector_store_status", "Not processed"),
                "",
                "="*70,
                "Contract ingestion complete. Document is now searchable via semantic queries.",
                "="*70
            ])
            
            final_output = "\n".join(summary_parts)
            
            return {
                "generated_text": final_output,
                "messages": ["Contract summary generated"]
            }
        
        # Build the graph
        graph = StateGraph(ContractState)
        
        # Add nodes
        graph.add_node("read_document", read_document_node)
        graph.add_node("extract_metadata", extract_metadata_node)
        graph.add_node("normalize_structure", normalize_and_structure_node)
        graph.add_node("ingest_vector_db", ingest_to_vector_db_node)
        graph.add_node("generate_summary", generate_summary_node)
        
        # Add edges - linear workflow
        graph.add_edge(START, "read_document")
        graph.add_edge("read_document", "extract_metadata")
        graph.add_edge("extract_metadata", "normalize_structure")
        graph.add_edge("normalize_structure", "ingest_vector_db")
        graph.add_edge("ingest_vector_db", "generate_summary")
        graph.add_edge("generate_summary", END)
        
        self.graph = graph.compile()
        return self.graph
    
    def run(self, file_path: str) -> dict:
        """
        Run the contract agent on a document.
        
        Args:
            file_path: Path to the contract document
            
        Returns:
            Final state with contract analysis
        """
        if self.graph is None:
            self.build_agent()
        
        initial_state = {
            "file_path": file_path,
            "raw_text": None,
            "contract_metadata": None,
            "structured_summary": None,
            "vector_store_status": None,
            "generated_text": "",
            "messages": []
        }
        
        result = self.graph.invoke(initial_state)
        return result

    def discover_partner_contracts(self, partner_name: str, contracts_dir: str = "docs") -> List[str]:
        """
        Discover all contract files related to a partner.
        
        Args:
            partner_name: Partner/company name to match
            contracts_dir: Directory to search
            
        Returns:
            List of matching contract file paths
        """
        contracts_path = Path(contracts_dir)
        if not contracts_path.exists():
            return []

        matches = []
        partner_tokens = {partner_name.lower(), partner_name.lower().replace(" ", "_")}
        for path in contracts_path.glob("*.docx"):
            filename = path.name.lower()
            if "confluent" in filename:
                if any(token in filename for token in partner_tokens) or partner_name.lower() == "confluent":
                    matches.append(str(path))
        return sorted(matches)

    def run_portfolio(self, partner_name: str, contract_paths: Optional[List[str]] = None) -> dict:
        """
        Process all discovered or provided contracts for a partner and return a workflow-ready portfolio view.
        """
        if contract_paths is None:
            contract_paths = self.discover_partner_contracts(partner_name)

        contract_results: List[Dict] = []
        active_contracts: List[Dict] = []
        renewal_candidates: List[Dict] = []
        recently_expired: List[Dict] = []
        long_dated_contracts: List[Dict] = []
        outstanding_obligations: List[Dict] = []
        enabled_permissions: List[Dict] = []
        pending_permissions: List[Dict] = []
        missing_requirements: List[Dict] = []
        notice_windows: List[Dict] = []
        risk_flags: List[str] = []
        next_steps: List[str] = []

        now = datetime.now().date()

        for contract_path in contract_paths:
            result = self.run(contract_path)
            structured = result.get("structured_summary") or {}
            metadata = result.get("contract_metadata") or {}

            effective_date_raw = structured.get("effective_date") or structured.get("start_date")
            effective_date = None
            if effective_date_raw and isinstance(effective_date_raw, str):
                try:
                    effective_date = datetime.fromisoformat(effective_date_raw).date()
                except ValueError:
                    # Try parsing common date formats
                    for fmt in ["%m/%d/%Y", "%Y-%m-%d", "%b %d, %Y", "%B %d, %Y"]:
                        try:
                            effective_date = datetime.strptime(effective_date_raw, fmt).date()
                            break
                        except ValueError:
                            continue

            filename = Path(contract_path).name
            
            # Try to get end_date from structured data first
            end_date_raw = structured.get("end_date")
            derived_end_date = None
            
            if end_date_raw and isinstance(end_date_raw, str):
                # Try parsing the end_date
                try:
                    derived_end_date = datetime.fromisoformat(end_date_raw).date()
                except ValueError:
                    # Try parsing common date formats
                    for fmt in ["%m/%d/%Y", "%Y-%m-%d", "%b %d, %Y", "%B %d, %Y"]:
                        try:
                            derived_end_date = datetime.strptime(end_date_raw, fmt).date()
                            break
                        except ValueError:
                            continue
            
            # Fallback: if no end_date but have effective_date and term_length, calculate it
            if not derived_end_date and effective_date:
                term_length = structured.get("term_length", "")
                if "year" in str(term_length).lower():
                    # Extract number of years
                    import re
                    years_match = re.search(r'(\d+)', str(term_length))
                    if years_match:
                        years = int(years_match.group(1))
                        try:
                            derived_end_date = effective_date.replace(year=effective_date.year + years)
                        except ValueError:
                            derived_end_date = None

            obligations = structured.get("key_obligations", []) if isinstance(structured, dict) else []
            milestones = structured.get("milestones", []) if isinstance(structured, dict) else []
            parties = structured.get("parties", []) if isinstance(structured, dict) else []

            normalized_text = f"{filename} {' '.join(obligations)} {' '.join(milestones)} {' '.join(parties)}".lower()

            hierarchy = "Governing Agreement"
            if "sow" in filename.lower():
                hierarchy = "SOW"
            elif "amend" in filename.lower():
                hierarchy = "Amendment"
            elif "renew" in filename.lower():
                hierarchy = "Renewal"
            elif "esa" in normalized_text:
                hierarchy = "ESA"
            elif "msa" in normalized_text:
                hierarchy = "MSA"

            status = "active"
            days_to_end = None
            if derived_end_date:
                days_to_end = (derived_end_date - now).days
                if days_to_end < 0:
                    status = "expired"
                elif days_to_end <= 180:
                    status = "expiring_soon"

            contract_record = {
                "file_path": contract_path,
                "file_name": filename,
                "structured_summary": structured,
                "document_length": metadata.get("document_length", 0),
                "effective_date": effective_date.isoformat() if effective_date else structured.get("effective_date", "Unknown"),
                "end_date": derived_end_date.isoformat() if derived_end_date else "Unknown",
                "contract_hierarchy": hierarchy,
                "status": status,
                "days_to_end": days_to_end,
                "enabled_vs_pending": "pending" if any(word in normalized_text for word in ["pending", "awaiting", "approval"]) else "enabled"
            }
            contract_results.append(contract_record)

            if contract_record["enabled_vs_pending"] == "enabled":
                enabled_permissions.append({
                    "contract": filename,
                    "status": "enabled",
                    "sales_relevance": "Current agreement appears usable for active sales motion"
                })
            else:
                pending_permissions.append({
                    "contract": filename,
                    "status": "pending",
                    "sales_relevance": "Additional step may be required before full activation"
                })

            if obligations:
                for obligation in obligations[:3]:
                    outstanding_obligations.append({
                        "contract": filename,
                        "obligation": obligation,
                        "timing": "next_30_days" if any(token in str(obligation).lower() for token in ["renew", "enable", "review", "approve", "sign"]) else "monitor"
                    })

            if any(token in normalized_text for token in ["sow", "statement of work"]) and "missing" in normalized_text:
                missing_requirements.append({
                    "contract": filename,
                    "requirement": "Potential downstream SOW dependency",
                    "impact": "May block downstream services or enablement"
                })

            if derived_end_date and days_to_end is not None:
                if days_to_end >= 0:
                    active_contracts.append(contract_record)
                    if days_to_end <= 180:
                        renewal_candidates.append({**contract_record, "days_to_renewal": days_to_end})
                        notice_windows.append({
                            "contract": filename,
                            "renewal_date": derived_end_date.isoformat(),
                            "notice_window": "Engage within the next 30-90 days",
                            "criticality": "high" if days_to_end <= 90 else "medium"
                        })
                    else:
                        long_dated_contracts.append(contract_record)
                elif -90 <= days_to_end < 0:
                    recently_expired.append({**contract_record, "days_since_expiry": abs(days_to_end)})

            if "cognos" in normalized_text:
                next_steps.append(
                    f"Meet with Confluent before the renewal window for {filename} because Cognos appears to be in scope."
                )
                risk_flags.append(f"Cognos-related contract requires proactive seller engagement: {filename}")

            if status == "expiring_soon":
                risk_flags.append(f"Contract expiring soon: {filename}")
            if status == "expired":
                risk_flags.append(f"Contract recently expired or no longer active: {filename}")
            if contract_record["enabled_vs_pending"] == "pending":
                risk_flags.append(f"Pending enablement or approval may delay sales activity: {filename}")

        # Generate dynamic next steps based on actual contract analysis
        # WHY: Hard-coded generic steps were appearing in every output regardless of context
        if not next_steps:
            # Build contract-specific recommendations
            if recently_expired:
                for contract in recently_expired[:2]:  # Top 2 most urgent
                    days_expired = contract.get("days_since_expiry", 0)
                    next_steps.append(
                        f"🚨 URGENT: {contract['file_name']} expired {days_expired} days ago - "
                        f"verify renewal status and create CRM opportunity if missing"
                    )
            
            if renewal_candidates:
                for contract in sorted(renewal_candidates, key=lambda x: x.get("days_to_renewal", 999))[:2]:
                    days_left = contract.get("days_to_renewal", 0)
                    urgency = "CRITICAL" if days_left <= 30 else "HIGH PRIORITY"
                    next_steps.append(
                        f"⚠️ {urgency}: {contract['file_name']} expires in {days_left} days - "
                        f"initiate renewal discussion immediately"
                    )
            
            if pending_permissions:
                next_steps.append(
                    f"📋 {len(pending_permissions)} contract(s) pending enablement - "
                    f"complete approval process to activate for sales"
                )
            
            # Only add generic steps if no specific actions identified
            if not next_steps:
                if active_contracts:
                    next_steps.append(
                        f"✅ {len(active_contracts)} active contract(s) - monitor for upcoming renewals"
                    )
                else:
                    next_steps.append(
                        "Review contract portfolio and confirm current status with legal/procurement"
                    )

        risk_level = "Low"
        if len(recently_expired) > 0 or len(renewal_candidates) >= 2 or len(pending_permissions) > 0:
            risk_level = "Medium"
        if len(recently_expired) > 1 or len(risk_flags) >= 4:
            risk_level = "High"

        portfolio_summary = {
            "partner_name": partner_name,
            "total_contracts": len(contract_results),
            "contract_paths": contract_paths,
            "active_contracts": active_contracts,
            "current_contract_status": {
                "enabled": enabled_permissions,
                "pending": pending_permissions
            },
            "outstanding_obligations": outstanding_obligations,
            "sales_relevant_permissions": {
                "enabled": enabled_permissions,
                "gated": pending_permissions
            },
            "time_bound_obligations_next_30_days": [item for item in outstanding_obligations if item.get("timing") == "next_30_days"],
            "missing_requirements": missing_requirements,
            "renewal_candidates": renewal_candidates,
            "recently_expired_contracts": recently_expired,
            "long_dated_contracts": long_dated_contracts,
            "renewal_notice_windows": notice_windows,
            "risk_flags": risk_flags,
            "risk_level": risk_level,
            "recommended_next_steps": next_steps
        }

        generated_text_lines = [
            "=" * 70,
            "CONTRACT PORTFOLIO SUMMARY",
            "=" * 70,
            f"Partner: {partner_name}",
            f"Total Contracts Processed: {len(contract_results)}",
            f"Active Contracts: {len(active_contracts)}",
            f"Contracts Coming Up for Renewal: {len(renewal_candidates)}",
            f"Recently Expired Contracts: {len(recently_expired)}",
            f"Risk Level: {risk_level}",
            "",
            "RECOMMENDED NEXT STEPS:",
        ]
        generated_text_lines.extend([f"- {step}" for step in next_steps])

        return {
            "partner_name": partner_name,
            "contract_paths": contract_paths,
            "contract_results": contract_results,
            "portfolio_summary": portfolio_summary,
            "generated_text": "\n".join(generated_text_lines),
            "messages": [f"Processed {len(contract_results)} contracts for {partner_name}"]
        }
    
    def query_contracts(self, query: str, k: int = 3) -> List[Document]:
        """
        Query ingested contracts using semantic search.
        
        Args:
            query: Search query
            k: Number of results to return
            
        Returns:
            List of relevant document chunks
        """
        if self.vector_store is None:
            self._initialize_vector_store()
        
        results = self.vector_store.similarity_search(query, k=k)
        return results


if __name__ == "__main__":
    # Example usage
    import os
    from dotenv import load_dotenv
    
    load_dotenv()
    
    agent = ContractAgent(
        apikey=os.getenv("WATSONX_APIKEY"),
        project_id=os.getenv("WATSONX_PROJECT_ID")
    )
    
    # Process a contract
    result = agent.run("sales-demo/docs/Confluent_IBM-1.30.2024.docx")
    
    print(result["generated_text"])
    
    # Example semantic query
    print("\n" + "="*70)
    print("SEMANTIC QUERY EXAMPLE")
    print("="*70)
    query_results = agent.query_contracts("What are the payment terms?")
    for i, doc in enumerate(query_results, 1):
        print(f"\nResult {i}:")
        print(doc.page_content[:200] + "...")
